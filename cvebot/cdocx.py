# import docx
import datetime
import os.path
from docx import Document

docx_directory_name = "cve_docx"
eng_to_rus_dict = {'id':'Идентификатор','affected':'Уязвимое ПО',"description":"Описание (оригинал)","description_ru":"Описание (перевод)","solution":"Устранение","link":"Ссылки","metrics":"Метрики","appear_time":"Дата регистрации","update_time":"Дата последнего обновления"}

def write_cve_to_docx(cve_dict):
    if not os.path.exists(docx_directory_name):
        os.mkdir(docx_directory_name)
        os.mkdir(os.path.join(docx_directory_name, datetime.datetime.now().strftime("%Y-%m-%d")))
    else:
        if not os.path.exists(os.path.join(docx_directory_name, datetime.datetime.now().strftime("%Y-%m-%d"))):
            os.mkdir(os.path.join(docx_directory_name, datetime.datetime.now().strftime("%Y-%m-%d")))
    """tmp_list = os.listdir(os.path.join(docx_directory_name, datetime.datetime.now().strftime("%Y-%m-%d")))
    sub_dir_last_file = False
    if tmp_list:
        sub_dir_last_file = [x for x in tmp_list if x and x.split("#")[0] == cve_dict['id']][-1]
    if sub_dir_last_file:
        print(sub_dir_last_file)
        document = Document()
        not_null_dict = {x: cve_dict[x] for x in cve_dict if cve_dict[x] and cve_dict[x] != ''}
        table = document.add_table(rows=len(not_null_dict), cols=2)
        i = 0
        for k in not_null_dict:
            cells_row = table.rows[i].cells
            cells_row[0].text = k
            cells_row[1].text = not_null_dict[k]
            i += 1
        document.save(os.path.join(docx_directory_name, datetime.datetime.now().strftime("%Y-%m-%d"),
                                   f"{cve_dict['id']}#{datetime.datetime.now().strftime('%Y-%m-%dT%H-%M-%S')}.docx"))
    else:"""
    document = Document()
    not_null_dict = {x: cve_dict[x] for x in cve_dict if cve_dict[x] and cve_dict[x] != ''}
    table = document.add_table(rows=len(not_null_dict), cols=2)
    table.style = 'Table Grid'
    i = 0
    for k in not_null_dict:
        cells_row = table.rows[i].cells
        cells_row[0].text = eng_to_rus_dict[k] if k in eng_to_rus_dict else k
        cells_row[0].paragraphs[0].runs[0].font.bold = True
        cells_row[1].text = not_null_dict[k]
        if k == 'id' or k == 'affected':
            cells_row[1].paragraphs[0].runs[0].font.bold = True
        i += 1
    file_string = os.path.join(docx_directory_name, datetime.datetime.now().strftime("%Y-%m-%d"),
                               f"{cve_dict['id']}#{datetime.datetime.now().strftime('%Y-%m-%dT%H-%M-%S')}.docx")
    document.save(file_string)
    return file_string
